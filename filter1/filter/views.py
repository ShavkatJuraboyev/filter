from django.shortcuts import render
from .models import Person
from django.db.models import Q
import openpyxl
from openpyxl.styles import Font, Alignment
from django.http import HttpResponse
from docx import Document
import django_filters
from django.template.loader import get_template
from django.views import View


def export_users_to_excel(request):
    # Foydalanuvchilarni o'qib olish
    all_persons = Person.objects.all()

    # Excel faylini yaratish
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = 'Persons'

    # Ma'lumotlarni yozish
    headers = ['First Name', 'Last Name', 'Age', 'Email', 'Status']

    header_font = Font(bold=True, color="FFFFFF")
    header_fill = openpyxl.styles.PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
    
    for col_num, header in enumerate(headers, 1):
        col_letter = openpyxl.utils.get_column_letter(col_num)
        cell = worksheet[f'{col_letter}1']
        cell.value = header
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    for row_num, person in enumerate(all_persons, 2):
        worksheet[f'A{row_num}'] = person.first_name
        worksheet[f'B{row_num}'] = person.last_name
        worksheet[f'C{row_num}'] = person.age
        worksheet[f'D{row_num}'] = person.email
        worksheet[f'E{row_num}'] = person.status

    # Excel faylini HttpResponse orqali qaytarish
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename=persons.xlsx'
    workbook.save(response)

    return response


def export_users_to_word(request):
    # Foydalanuvchilarni o'qib olish
    all_persons = Person.objects.all()

    # Word faylini yaratish
    document = Document()
    document.add_heading('User List', level=1)

    # Ma'lumotlarni chiroyliroq qatorma-qator yozish
    for person in all_persons:
        document.add_paragraph(f"Name: {person.first_name} {person.last_name}")
        document.add_paragraph(f"Age: {person.age}")
        document.add_paragraph(f"Email: {person.email}")
        document.add_paragraph(f"Status: {person.status}")
        document.add_paragraph("-" * 30)  # Har bir foydalanuvchi uchun ayiruvchi chiziq

    # Word faylini HttpResponse orqali qaytarish
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=persons.docx'
    document.save(response)

    return response


def person_list(request):
    # Barcha foydalanuvchilarni olish
    all_persons = Person.objects.all()

    # Foydalanuvchilarni 18 yoshdan katta bo'lganlarni olish
    adult_persons = Person.objects.filter(age__gt=18)

    return render(request, 'status.html', {'all_persons': all_persons, 'adult_persons': adult_persons})


def person_list1(request):
    # Barcha foydalanuvchilar ro'yxati
    all_persons = Person.objects.all()

    # Qidirish parametrlarini olishish
    search_query = request.GET.get('search', '')
    status_filter = request.GET.get('status', 'all')
    age_filter = request.GET.get('age', 'all')

    # Qidiruvni amalga oshirish
    if search_query:
        all_persons = all_persons.filter(
        Q(first_name__icontains=search_query) |
        Q(last_name__icontains=search_query) |
        Q(age__icontains=search_query) |
        Q(email__icontains=search_query) |
        Q(status__icontains=search_query)
    )

    if status_filter != 'all':
        all_persons = all_persons.filter(status=status_filter)

    if age_filter != 'all':
        all_persons = all_persons.filter(age=age_filter)

    context = {
        'all_persons': all_persons,
        'search_query': search_query,
        'status_filter': status_filter,
        'age_filter': age_filter,
    }

    return render(request, 'django.html', context)
