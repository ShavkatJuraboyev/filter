from django.shortcuts import render
from .models import Person, Tuman
from django.http import HttpResponse
import openpyxl
from docx import Document

def person(request):
    persons = Person.objects.all()

    first_name = request.GET.get('first_name', '')
    last_name = request.GET.get('last_name', '')
    age = request.GET.get('age', '')
    tuman_id = request.GET.get('tuman', '')
    status = request.GET.get('status', '')

    if first_name:
        persons = persons.filter(first_name__icontains=first_name)
    if last_name:
        persons = persons.filter(last_name__icontains=last_name)
    if age:
        persons = persons.filter(age=age)
    if tuman_id:
        persons = persons.filter(tuman_id=tuman_id)
    if status:
        persons = persons.filter(status=status)

    tumans = Tuman.objects.all()
    status_choices = Person.STATUS_CHOICES

    ctx = {'persons': persons, 'tumans': tumans, 'status_choices': status_choices}

    # Export to Excel
    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.append(['First Name', 'Last Name', 'Age', 'Tuman', 'Status'])

    for person in persons:
        # Convert the Tuman object to a string or extract a specific field
        tuman_value = str(person.tuman)  # or person.tuman.tuman_name if 'tuman_name' is the field
        worksheet.append([person.first_name, person.last_name, person.age, tuman_value, person.status])


    excel_response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    excel_response['Content-Disposition'] = 'attachment; filename=persons.xlsx'
    workbook.save(excel_response)

    # Export to Word
    document = Document()
    document.add_heading('Persons Report', 0)
    
    for person in persons:
        document.add_paragraph(f"{person.first_name} {person.last_name}, Age: {person.age}, Tuman: {person.tuman}, Status: {person.status}")

    word_response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    word_response['Content-Disposition'] = 'attachment; filename=persons.docx'
    document.save(word_response)

    return render(request, 'person/person.html', ctx)
